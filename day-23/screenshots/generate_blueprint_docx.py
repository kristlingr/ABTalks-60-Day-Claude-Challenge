"""Generate Customer & MVP Blueprint as a Word document."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).parent / "Customer_MVP_Blueprint.docx"

NAVY = RGBColor(15, 52, 96)
ACCENT = RGBColor(233, 69, 96)


def set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement

    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level <= 2 else RGBColor(22, 33, 62)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr_cells[i], "0F3460")

    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = val
            for p in row_cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
        if r_idx % 2 == 1:
            for cell in row_cells:
                set_cell_shading(cell, "F4F7FB")

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(10)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title.add_run("Customer & MVP Blueprint")
    t_run.bold = True
    t_run.font.size = Pt(22)
    t_run.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = sub.add_run("NightShift Kitchen | Pune, India | June 2026")
    s_run.font.size = Pt(11)
    s_run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

    # Executive Summary
    add_heading(doc, "Executive Summary", 1)
    doc.add_paragraph(
        "NightShift Kitchen targets a real, recurring pain: IT night support workers in Pune "
        "cannot access hygienic, affordable meals between 10 PM–5 AM. The founder has 2 years of "
        "lived experience and qualitative on-shift validation. India's midnight food economy "
        "(15–20% of delivery orders, 11 PM–2 AM) and cloud kitchen growth (12–17% CAGR) support timing."
    )
    p = doc.add_paragraph()
    p.add_run("Strategic wedge: ").bold = True
    p.add_run('Own "trusted hygienic meals for IT night crews" — not generic late-night delivery.')

    add_table(
        doc,
        ["Factor", "Status"],
        [
            ["Problem severity", "✅ High (7.4/10)"],
            ["Market size (SAM)", "✅ ₹850–1,200 Cr"],
            ["Founder–market fit", "✅ Strong (7.4/10)"],
            ["Willingness to pay", "❌ Unvalidated"],
            ["Unit economics", "⚠️ Tight (15–25% gross margin)"],
            ["Competitive moat", "⚠️ Moderate (niche + trust)"],
        ],
        [1.8, 4.5],
    )
    doc.add_paragraph(
        "Recommendation: Lean 30-day pilot in Hinjewadi. Do not lease a full kitchen or quit day job "
        "until 50 pre-orders and ₹140 AOV are proven."
    )

    # ICP
    add_heading(doc, "Ideal Customer Profile", 1)
    add_table(
        doc,
        ["Attribute", "Tier 1 ICP"],
        [
            ["Role", "L1/L2 IT Support, Helpdesk Analyst, NOC Engineer, Application Support, BPO Agent"],
            ["Industry", "IT Services, BPO/KPO, GCC, 24/7 product support"],
            ["Company size", "200–10,000 employees"],
            ["Age / income", "22–35 yrs; ₹25K–60K/month + ₹2K–4K shift allowance"],
            ["Location", "Pune — Hinjewadi (P0), Magarpatta, Kharadi, Balewadi"],
            ["Shift", "8 PM – 6 AM IST (US/UK client alignment)"],
            ["Behavior", "Orders via apps/WhatsApp; currently Swiggy, tiffin, Maggi, skips meals"],
            ["Budget", "₹100–180/meal; 8–15 external meals/month"],
            ["Pain", "No hygienic options post-10 PM; unhealthy eating cycle"],
        ],
        [1.6, 4.7],
    )

    # Buyer Persona
    add_heading(doc, "Buyer Persona: Rahul — The Night Shift Warrior", 1)
    add_table(
        doc,
        ["Field", "Detail"],
        [
            ["Age / role", "27, L2 Application Support Engineer"],
            ["Company", "Mid-size IT firm, Hinjewadi Phase 2"],
            ["Shift / pay", "9 PM–6 AM; ₹42K + ₹3,500 night allowance"],
            ["Living", "Shared PG in Wakad, 15 min from office"],
        ],
        [1.5, 4.8],
    )
    quote = doc.add_paragraph()
    quote.add_run(
        '"I don\'t need fancy food. I need something clean, filling, and under ₹150 that I can '
        'order at 1 AM without gambling on food poisoning."'
    ).italic = True

    add_table(
        doc,
        ["Goals", "Frustrations"],
        [
            ["Eat healthier without cooking", "Swiggy at 2 AM feels oily and risky"],
            ["Save money on food", "Biryani nightly = ₹6,000/month"],
            ["Stay alert on shift", "Maggi causes 4 AM energy crash"],
            ["Simple, predictable meals", "500-menu apps are overwhelming"],
        ],
        [2.4, 3.9],
    )

    # Pain Points
    add_heading(doc, "Top 10 Customer Pain Points", 1)
    add_table(
        doc,
        ["#", "Pain Point", "Severity", "Frequency", "Workaround"],
        [
            ["1", "No hygienic food after 10 PM", "Critical", "Daily", "Swiggy gamble / skip meal"],
            ["2", "Unhealthy eating → weight/health", "Critical", "Weekly", "Ignore it"],
            ["3", "High delivery cost + fees", "High", "Per order", "Group orders"],
            ["4", "Long delivery wait (45+ min)", "High", "2–3×/week", "Order early, eat cold"],
            ["5", "Limited night menu", "Medium", "Daily", "Rotate same 3 restaurants"],
            ["6", "Office canteen closes early", "High", "Daily", "Pack boring tiffin"],
            ["7", "Irregular meal timing / crashes", "Medium", "Daily", "Energy drinks"],
            ["8", "Food safety anxiety", "High", "Per order", "Stick to known places"],
            ["9", "Tiffin fatigue by Day 3", "Medium", "Weekly", "Switch to junk food"],
            ["10", "No shift-aligned meal bundles", "Medium", "Daily", "Snack randomly"],
        ],
        [0.3, 2.0, 0.7, 0.7, 2.0],
    )

    # Customer Journey
    add_heading(doc, "Customer Journey", 1)
    add_table(
        doc,
        ["Stage", "Action", "Touchpoint", "Opportunity"],
        [
            ["Awareness", "Complains about food on shift", "Colleague WOM, WhatsApp", "Campus posters, IT park flyers"],
            ["Consideration", 'Searches "food delivery Hinjewadi night"', "Instagram, Google, Swiggy", "SEO + aggregator listing"],
            ["Purchase", "First order at discount", "WhatsApp / app", "₹99 first meal; hygiene packaging"],
            ["Activation", "2nd order within 7 days", "WhatsApp reorder link", "10:30 PM reminder push"],
            ["Retention", "3+ orders/week", "Subscription / favorites", "Meal plan subscription"],
            ["Advocacy", "Refers team", "Referral code", "Refer 3, get 5 free meals"],
            ["Churn risk", "Stops ordering", "—", "Weekly menu rotation"],
        ],
        [1.0, 1.5, 1.3, 2.5],
    )

    # Objections
    add_heading(doc, "Key Customer Objections", 1)
    add_table(
        doc,
        ["Objection", "Likelihood", "Response"],
        [
            ["Swiggy already delivers at night", "HIGH", "Built for your shift — hygienic, affordable, always available"],
            ["I bring tiffin from home", "HIGH", "Backup when tiffin gets boring, not a replacement"],
            ["₹150 is too much", "MEDIUM", "₹99 snack box; ₹2,999/month (20 meals) subscription"],
            ["I don't trust cloud kitchens", "MEDIUM", "FSSAI license, kitchen transparency, founder story"],
            ["Delivery takes too long", "MEDIUM", "30-min guarantee in Hinjewadi pilot zone"],
            ["I'll try later", "HIGH", "First order free (up to ₹120); 1-tap WhatsApp reorder"],
        ],
        [1.8, 0.8, 3.7],
    )

    # Buying Triggers
    add_heading(doc, "Key Buying Triggers", 1)
    add_table(
        doc,
        ["Trigger", "When", "Marketing Response"],
        [
            ["Hunger spike", "11 PM–1 AM", "Push/WhatsApp at 10:45 PM"],
            ["Bad food experience", "After food poisoning/stale order", '"Never gamble again" trust messaging'],
            ["Payday", "1st–5th of month", "Subscription / meal pack offers"],
            ["New night-shift joiner", "Onboarding week", "Campus flyer + 50% off first order"],
            ["Colleague referral", "Team orders together", "Free delivery for 3+ orders"],
            ["Health scare", "Weight gain, doctor visit", '"Night shift meal plan" content'],
            ["Canteen closure", "Office policy change", "B2B outreach to team leads"],
        ],
        [1.4, 1.5, 3.4],
    )

    # MVP Recommendation
    add_heading(doc, "MVP Recommendation", 1)
    add_heading(doc, "What to Build First", 2)
    add_table(
        doc,
        ["Priority", "Build", "Why"],
        [
            ["1", "8-item curated night menu", "Low COGS, fast ops, matches persona need"],
            ["2", "WhatsApp Business ordering + pre-order page", "Zero app cost; matches behavior"],
            ["3", "FSSAI-licensed shared kitchen in Hinjewadi/Wakad", "Compliance + low capex"],
            ["4", "Geo-fenced delivery (Hinjewadi Phase 1–3)", "Predictable routes, 30-min SLA"],
            ["5", "Hygiene trust layer (FSSAI, packaging, founder video)", "Core differentiator vs aggregators"],
            ["6", "Swiggy/Zomato listing (discovery only)", "Drive trials; push direct reorder"],
            ["7", "₹99 launch offer + meal subscription", "Prove WTP and repeat behavior"],
        ],
        [0.6, 2.2, 3.5],
    )

    add_heading(doc, "What NOT to Build (Yet)", 2)
    add_bullets(
        doc,
        [
            "Custom mobile app",
            "Multi-city expansion",
            "Wide cuisine variety (biryani, Chinese, etc.)",
            "Own kitchen lease / full-time team",
            "Paid performance marketing",
            "B2B corporate contracts",
            "Live kitchen cam / fancy branding",
            "24/7 operations (start 4 nights/week, 10 PM–3 AM)",
        ],
    )

    add_heading(doc, "Success Metrics (30–60 Days)", 2)
    add_table(
        doc,
        ["Metric", "Target", "Kill Signal"],
        [
            ["Pre-orders", "≥50", "<30"],
            ["First-month orders", "≥200", "<100 in 60 days"],
            ["2nd-order rate (7 days)", "≥40%", "<25%"],
            ["Repeat (3×/week)", "25% of actives", "<15%"],
            ["AOV", "₹129–140", "Must drop below ₹99 to convert"],
            ["Gross margin", "≥15%", "<10%"],
            ["Customer satisfaction", "≥7/10", "Any food safety incident"],
            ["Nightly orders (Week 4)", "≥15/night", "Consistent <8/night"],
        ],
        [2.0, 1.5, 2.8],
    )

    # MoSCoW
    add_heading(doc, "MoSCoW Prioritization", 1)
    add_table(
        doc,
        ["Must Have", "Should Have", "Could Have", "Won't Have (MVP)"],
        [
            ["8-item hygienic menu", "Swiggy/Zomato listing", "Weekly menu rotation", "Custom app"],
            ["FSSAI shared kitchen", "2–3 dedicated night riders", "Referral program", "Multi-city ops"],
            ["WhatsApp ordering", "₹99 launch + subscription", "Instagram content", "B2B corporate sales"],
            ["Hinjewadi-only zone", "Campus flyers + WOM", "Google Maps / SEO", "Full kitchen lease"],
            ["COGS <35% per item", "Pre-order landing page", "Pickup locker at IT gate", "Paid ads"],
            ["Part-time chef/manager", "Customer feedback loop", "Health meal plan", "24/7 full coverage"],
            ["Daily ops P&L tracking", "30-min delivery guarantee", "Employer wellness pitch", "Wide cuisine menu"],
        ],
        [1.5, 1.5, 1.5, 1.5],
    )

    # Pricing
    add_heading(doc, "Pricing Hypothesis", 1)
    add_table(
        doc,
        ["Tier", "Price", "Predicted Conversion", "Role"],
        [
            ["Snack box", "₹79", "70%", "Volume driver, low margin"],
            ["Standard meal", "₹129", "55%", "Sweet spot — anchor here"],
            ["Premium meal", "₹179", "30%", "Weekend / treat orders"],
            ["Monthly pack", "₹2,999 (20 meals)", "15%", "Retention / loyalty play"],
        ],
        [1.2, 1.0, 1.3, 2.8],
    )
    doc.add_paragraph(
        "Unit economics target: Food COGS 32–38% | Delivery 15–22% | Gross margin 15–25% | "
        "Break-even: 40–60 orders/day at ₹140 AOV."
    )

    # Risks
    add_heading(doc, "Top 5 Risks", 1)
    add_table(
        doc,
        ["#", "Risk", "Score", "Mitigation"],
        [
            ["1", "Aggregator dominance (Swiggy/Zomato)", "9", "Direct ordering; niche brand; campus marketing"],
            ["2", "Thin margins (food + delivery + commission)", "9", "Commission-free direct orders; COGS <35%"],
            ["3", "Low order density at night", "7", "Geo-fence IT parks; batch cooking"],
            ["4", "Founder bandwidth (side business)", "7", "Hire kitchen manager; 8-item menu max"],
            ["5", "FSSAI / hygiene failure", "7", "Licensed kitchen; audits; insurance"],
        ],
        [0.3, 2.2, 0.5, 3.3],
    )

    # 30-Day Plan
    add_heading(doc, "30-Day MVP Plan", 1)
    add_table(
        doc,
        ["Week", "Focus", "Key Actions", "Success Gate"],
        [
            ["W1: Validate", "Survey + brand", "50-person survey; brand + 8-item menu; shortlist kitchens; WhatsApp", "≥40 responses; ≥60% order 3×/week"],
            ["W2: Pre-sell", "Kitchen + pre-orders", "Sign LOI; COGS sheet; pre-order page; 30 packs at ₹99; hire rider", "Kitchen secured; COGS <35%; ≥30 pre-orders"],
            ["W3: Soft launch", "First deliveries", "FSSAI app; test batches; deliver pre-orders; Swiggy/Zomato; flyers", "30 meals, 0 complaints; ≥50 orders"],
            ["W4: Measure", "Ops + decision", "Nightly ops; track P&L; interview 10 customers; repeat rate", "≥15 orders/night; margin ≥15%; ≥7/10 satisfaction"],
        ],
        [0.9, 1.0, 2.5, 2.0],
    )

    # Founder Action Sheet
    add_heading(doc, "Founder Action Sheet — Top 10 Next Actions", 1)
    add_table(
        doc,
        ["#", "Action", "Deadline", "Owner"],
        [
            ["1", "Launch 10-question Google Form survey to 50 night-shift contacts", "Day 3", "Founder"],
            ["2", "Analyze WTP, preferred price, meal times from survey", "Day 4", "Founder"],
            ["3", "Finalize brand name, positioning, 8-item menu", "Day 5", "Founder + Chef"],
            ["4", "Shortlist & visit 3 FSSAI shared kitchens (Hinjewadi/Wakad)", "Day 6–9", "Founder"],
            ["5", "Sign kitchen LOI; design menu with chef; COGS <35%", "Day 10", "Founder + Chef"],
            ["6", "Create WhatsApp Business + pre-order landing page", "Day 11", "Founder"],
            ["7", "Pre-sell 30 meal packs at ₹99 launch price", "Day 12–13", "Founder"],
            ["8", "Hire 1 part-time delivery rider (10 PM–3 AM routes)", "Day 14", "Founder"],
            ["9", "Soft launch: deliver pre-orders; list on Swiggy/Zomato", "Day 17–18", "Founder + Ops"],
            ["10", "Run nightly ops Week 4; calculate P&L; interview 10 customers", "Day 22–27", "Founder"],
        ],
        [0.3, 3.2, 0.8, 0.9],
    )

    # Scores
    add_heading(doc, "Scores (0–100)", 1)
    add_table(
        doc,
        ["Dimension", "Score", "Rationale"],
        [
            ["Customer Clarity", "82", "Sharp ICP; persona well-defined; geo-clustered"],
            ["Problem Severity", "78", "Daily high-intensity pain; founder lived it; WTP unproven"],
            ["PMF Potential", "68", "Real gap + founder access; low switching costs; thin margins"],
            ["MVP Readiness", "62", "Clear 30-day plan; ops partner gap; no pre-orders yet"],
        ],
        [1.4, 0.7, 3.2],
    )

    # Final Verdict
    add_heading(doc, "Final Verdict", 1)
    verdict = doc.add_paragraph()
    verdict.alignment = WD_ALIGN_PARAGRAPH.CENTER
    v_run = verdict.add_run("🟡 Promising but Unvalidated")
    v_run.bold = True
    v_run.font.size = Pt(14)
    v_run.font.color.rgb = ACCENT

    doc.add_paragraph(
        "NightShift Kitchen addresses a validated, high-frequency problem with strong founder–market fit "
        "and a defensible niche wedge (IT night crews + hygiene trust). Market timing and geo-density in "
        "Pune IT parks are favorable."
    )
    doc.add_paragraph(
        "However, willingness to pay is unproven, unit economics are tight, and competitive intensity "
        "is high. The venture is ready to test — not scale."
    )
    doc.add_paragraph(
        "Next gate: 50 pre-orders + ₹140 AOV acceptance + 40% repeat rate within 60 days. "
        "Pass → pursue subscription model and B2B. Fail → pivot to B2B corporate catering or subscription tiffin."
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = footer.add_run("Source: Startup Validation Report v2 (June 24, 2026) | Confidential — Internal Use")
    f_run.font.size = Pt(8)
    f_run.font.color.rgb = RGBColor(120, 120, 120)

    doc.save(OUT)
    print(f"Created: {OUT} ({OUT.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    build()
